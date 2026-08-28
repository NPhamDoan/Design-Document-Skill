"""Chuyển tài liệu thiết kế Markdown -> Word (.docx) bằng python-docx.

Dùng khi tài liệu ĐÃ ở dạng Markdown (heading + bảng + ảnh) và cần bản .docx,
không có/không muốn cài pandoc. Hoạt động đa nền tảng, chỉ cần python-docx (+ Pillow
để co ảnh đúng tỉ lệ).

Hỗ trợ: heading (#..######), đoạn văn, **đậm**, `code`, danh sách (-, *, 1.),
bảng markdown (| ... |), ảnh (![alt](path)), blockquote (>), code block (```),
đường kẻ ngang (---). Ảnh nhúng và co theo bề rộng/cao trang, đường dẫn tính
tương đối theo file .md.

Cài phụ thuộc:
    pip install python-docx Pillow

Chạy:
    python md-to-docx.py <input.md> <output.docx>
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

try:
    from PIL import Image
    HAS_PIL = True
except Exception:
    HAS_PIL = False

# Bề rộng/cao tối đa cho ảnh (trong lề trang A4, lề 2.54cm)
MAX_IMG_W_CM = 15.5
MAX_IMG_H_CM = 20.0
BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"


def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:fill"): color, qn("w:val"): "clear", qn("w:color"): "auto",
    })
    tcPr.append(shd)


def set_para_shading(paragraph, color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:fill"): color, qn("w:val"): "clear", qn("w:color"): "auto",
    })
    pPr.append(shd)


# ---- Inline: **đậm**, `code` ----
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(paragraph, text):
    text = text.replace("&nbsp;", " ")
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO_FONT
        else:
            paragraph.add_run(part)


def img_size_cm(path):
    """Trả (width_cm, height_cm) đã co vừa MAX box, giữ tỉ lệ."""
    if not HAS_PIL:
        return MAX_IMG_W_CM, None
    with Image.open(path) as im:
        w_px, h_px = im.size
    w_cm = w_px / 96 * 2.54
    h_cm = h_px / 96 * 2.54
    if not w_cm or not h_cm:
        return MAX_IMG_W_CM, None
    scale = min(MAX_IMG_W_CM / w_cm, MAX_IMG_H_CM / h_cm)
    return w_cm * scale, h_cm * scale


class MdToDocx:
    def __init__(self, md_path: Path):
        self.md_path = md_path
        self.base = md_path.parent
        self.doc = Document()
        self._setup()

    def _setup(self):
        for s in self.doc.sections:
            s.top_margin = s.bottom_margin = Cm(2.54)
            s.left_margin = s.right_margin = Cm(2.54)
        normal = self.doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(12)

    def add_heading(self, text, level):
        self.doc.add_heading(text, level=min(level, 4))

    def add_paragraph(self, text):
        add_inline(self.doc.add_paragraph(), text)

    def add_bullet(self, text, level=0):
        para = self.doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Cm(0.5 + 0.5 * level)
        add_inline(para, text)

    def add_number(self, text):
        add_inline(self.doc.add_paragraph(style="List Number"), text)

    def add_quote(self, text):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def add_code_block(self, lines):
        para = self.doc.add_paragraph()
        set_para_shading(para, "F2F2F2")
        for i, ln in enumerate(lines):
            run = para.add_run(ln)
            run.font.name = MONO_FONT
            run.font.size = Pt(10)
            if i < len(lines) - 1:
                run.add_break()

    def add_image(self, alt, rel_path):
        img_path = (self.base / rel_path).resolve()
        if not img_path.exists():
            self.add_paragraph(f"[Thiếu ảnh: {rel_path}]")
            return
        w_cm, h_cm = img_size_cm(img_path)
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        if h_cm:
            run.add_picture(str(img_path), width=Cm(w_cm), height=Cm(h_cm))
        else:
            run.add_picture(str(img_path), width=Cm(w_cm))
        if alt:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(alt)
            r.italic = True
            r.font.size = Pt(10)

    def add_table(self, header, rows):
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(header))
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(header):
            cell = tbl.rows[0].cells[i]
            cell.text = ""
            add_inline(cell.paragraphs[0], h)
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10.5)
            set_cell_shading(cell, "D9E2F3")
        for r_idx, row in enumerate(rows):
            for c_idx in range(len(header)):
                val = row[c_idx] if c_idx < len(row) else ""
                cell = tbl.rows[r_idx + 1].cells[c_idx]
                cell.text = ""
                add_inline(cell.paragraphs[0], val)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10.5)
        self.doc.add_paragraph()


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_IMG_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")


def split_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def convert(md_path: Path, out_path: Path):
    conv = MdToDocx(md_path)
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            conv.add_code_block(block)
            continue

        if not stripped:
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            i += 1
            continue

        m = _HEADING_RE.match(line)
        if m:
            conv.add_heading(m.group(2).strip(), max(len(m.group(1)) - 1, 0))
            i += 1
            continue

        m = _IMG_RE.match(stripped)
        if m:
            conv.add_image(m.group(1).strip(), m.group(2).strip())
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1]):
            header = split_table_row(lines[i])
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            conv.add_table(header, rows)
            continue

        if stripped.startswith(">"):
            conv.add_quote(stripped.lstrip(">").strip())
            i += 1
            continue

        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            conv.add_bullet(m.group(2).strip(), len(m.group(1)) // 2)
            i += 1
            continue

        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            conv.add_number(m.group(1).strip())
            i += 1
            continue

        conv.add_paragraph(stripped)
        i += 1

    conv.doc.save(str(out_path))
    print(f"Da tao: {out_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Cach dung: python md-to-docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
